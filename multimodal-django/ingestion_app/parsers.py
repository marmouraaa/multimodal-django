# ingestion_app/parsers.py
import os
import json
import logging
from pathlib import Path
from datetime import datetime
import PyPDF2
from docx import Document
logger = logging.getLogger(__name__)


def normalize_text(text: str) -> str:
    """Normalise le texte extrait"""
    if not text:
        return ''
    import re
    text = text.replace('\x00', '')
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'\n{3,}', '\n\n', text)
    lines = [line.rstrip() for line in text.split('\n')]
    text = '\n'.join(lines)
    return text.strip()


def parse_pdf(file_path: str) -> dict:
    """Extrait le texte d'un fichier PDF"""
    try:
     
        result = {'text': '', 'page_count': 0, 'metadata': {}, 'error': None}
        
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            result['page_count'] = len(reader.pages)
            
            if reader.metadata:
                meta = reader.metadata
                result['metadata'] = {
                    'author': meta.get('/Author', ''),
                    'title': meta.get('/Title', ''),
                    'subject': meta.get('/Subject', ''),
                }
            
            pages_text = []
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(f"--- Page {i + 1} ---\n{page_text.strip()}")
                except Exception as e:
                    logger.warning(f"Erreur page {i + 1}: {e}")
            
            result['text'] = normalize_text('\n\n'.join(pages_text))
        return result
    except Exception as e:
        logger.error(f"Erreur parsing PDF: {e}")
        return {'text': '', 'page_count': 0, 'metadata': {}, 'error': str(e)}


def parse_docx(file_path: str) -> dict:
    """Extrait le texte d'un fichier Word"""
    try:
        
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = '\n'.join(paragraphs)
        
        # Extraire les tableaux
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(' | '.join(cells))
            if rows:
                full_text += f"\n\n[Tableau {i+1}]\n" + '\n'.join(rows)
        
        return {'text': normalize_text(full_text), 'error': None}
    except Exception as e:
        logger.error(f"Erreur parsing DOCX: {e}")
        return {'text': '', 'error': str(e)}


def parse_txt(file_path: str) -> dict:
    """Lit un fichier texte brut"""
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return {'text': normalize_text(f.read()), 'encoding': encoding, 'error': None}
        except UnicodeDecodeError:
            continue
    return {'text': '', 'error': "Impossible de décoder le fichier"}


def parse_image(file_path: str) -> dict:
    """Charge et normalise une image"""
    try:
        from PIL import Image
        img = Image.open(file_path)
        original_size = img.size
        
        description = f"Image de {original_size[0]}x{original_size[1]} pixels, fichier: {Path(file_path).name}"
        img.close()
        
        return {'description': description, 'width': original_size[0], 'height': original_size[1], 'error': None}
    except Exception as e:
        logger.error(f"Erreur parsing image: {e}")
        return {'description': '', 'error': str(e)}


def parse_csv(file_path: str) -> dict:
    """Lit un fichier CSV"""
    try:
        import pandas as pd
        df = pd.read_csv(file_path, encoding='utf-8')
        return _dataframe_to_result(df, file_path)
    except Exception as e:
        logger.error(f"Erreur parsing CSV: {e}")
        return {'text': '', 'error': str(e)}


def parse_excel(file_path: str) -> dict:
    """Lit un fichier Excel"""
    try:
        import pandas as pd
        df = pd.read_excel(file_path)
        return _dataframe_to_result(df, file_path)
    except Exception as e:
        logger.error(f"Erreur parsing Excel: {e}")
        return {'text': '', 'error': str(e)}


def parse_json(file_path: str) -> dict:
    """Lit un fichier JSON"""
    try:
        import pandas as pd
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        if isinstance(data, list) and data and isinstance(data[0], dict):
            df = pd.DataFrame(data)
            return _dataframe_to_result(df, file_path)
        elif isinstance(data, dict):
            text_lines = [f"Structure JSON :"]
            for key, value in data.items():
                text_lines.append(f"  {key}: {str(value)[:100]}")
            return {'text': normalize_text('\n'.join(text_lines)), 'error': None}
        else:
            return {'text': normalize_text(str(data)), 'error': None}
    except Exception as e:
        logger.error(f"Erreur parsing JSON: {e}")
        return {'text': '', 'error': str(e)}


def _dataframe_to_result(df, file_path: str) -> dict:
    """Convertit un DataFrame en texte descriptif"""
    import pandas as pd
    
    text_lines = [
        f"Fichier: {Path(file_path).name}",
        f"Dimensions: {len(df)} lignes x {len(df.columns)} colonnes",
        f"Colonnes: {', '.join(str(c) for c in df.columns)}",
        ""
    ]
    
    # Statistiques pour colonnes numériques
    numeric_cols = df.select_dtypes(include=['number']).columns
    if len(numeric_cols) > 0:
        text_lines.append("=== Statistiques ===")
        for col in numeric_cols:
            col_data = df[col].dropna()
            if len(col_data) > 0:
                text_lines.append(
                    f"{col}: min={col_data.min():.2f}, max={col_data.max():.2f}, "
                    f"moyenne={col_data.mean():.2f}"
                )
        text_lines.append("")
    
    # Aperçu des données
    text_lines.append("=== Aperçu (10 premières lignes) ===")
    text_lines.append(df.head(10).to_string(index=False))
    
    return {'text': normalize_text('\n'.join(text_lines)), 'error': None}


def run_parsing_pipeline(file_path: str, original_name: str) -> dict:
    """Pipeline complet d'ingestion + parsing"""
    logger.info(f"Début pipeline pour: {original_name}")
    
    ext = Path(file_path).suffix.lower()
    result = {'success': True, 'errors': [], 'warnings': [], 'extracted_text': '', 'extra': {}}
    
    # Parser selon l'extension
    if ext == '.pdf':
        parsed = parse_pdf(file_path)
        result['modality'] = 'document'
        result['mime_type'] = 'application/pdf'
    elif ext in ('.docx', '.doc'):
        parsed = parse_docx(file_path)
        result['modality'] = 'document'
    elif ext in ('.txt', '.md'):
        parsed = parse_txt(file_path)
        result['modality'] = 'document'
    elif ext in ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'):
        parsed = parse_image(file_path)
        result['modality'] = 'image'
        result['extracted_text'] = parsed.get('description', '')
        result['extra']['width'] = parsed.get('width')
        result['extra']['height'] = parsed.get('height')
    elif ext in ('.xlsx', '.xls'):
        parsed = parse_excel(file_path)
        result['modality'] = 'structured'
    elif ext == '.json':
        parsed = parse_json(file_path)
        result['modality'] = 'structured'
    elif ext in ('.csv', '.tsv'):
        parsed = parse_csv(file_path)
        result['modality'] = 'structured'
    else:
        parsed = {'error': f'Extension non supportée: {ext}', 'text': ''}
        result['success'] = False
        result['errors'].append(f"Extension non supportée: {ext}")
    
    if parsed.get('error'):
        result['errors'].append(parsed['error'])
        result['success'] = False
    else:
        if result['modality'] != 'image':
            result['extracted_text'] = parsed.get('text', '')
    
    logger.info(f"Pipeline terminé: modality={result.get('modality')}, chars={len(result.get('extracted_text', ''))}")
    return result